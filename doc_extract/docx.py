import os
from docx import Document
import docx2txt
from doc_extract.images_descriptions import ImageExtractions

class DocxExtracting:
    def __init__(self, path_to_docx):
        """
        Инициализирует класс с путём к файлу .docx.
        :param path_to_docx: Путь к файлу .docx
        """
        self.path = path_to_docx  # Сохраняем путь к файлу .docx

    def text_extract(self):
        """
        Извлекает текст из параграфов и таблиц документа.
        :return: Словарь, где ключи — номера параграфов и таблиц, а значения — их содержимое
        """
        doc = Document(self.path)  # Открываем документ по указанному пути
        content = {}  # Создаём словарь для хранения содержимого документа

        # Извлекаем текст из всех параграфов
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            content[f'Параграф{i}'] = paragraph.text  # Сохраняем текст параграфа с ключом "ПараграфN"

        # Извлекаем данные из всех таблиц
        for table_index, table in enumerate(doc.tables, start=1):
            headers = [cell.text.strip() for cell in table.rows[0].cells]  # Извлекаем заголовки из первой строки таблицы
            data = []  # Создаём список для хранения данных таблицы
            for row in table.rows[1:]:  # Проходим по всем строкам таблицы, кроме заголовков
                row_data = [cell.text.strip() for cell in row.cells]  # Извлекаем данные из ячеек строки
                data.append(row_data)  # Добавляем данные строки в список
            content[f'Таблица{table_index}'] = {'headers': headers, 'data': data}  # Сохраняем данные таблицы с ключом "ТаблицаN"

        return content  # Возвращаем словарь с содержимым документа

    def det_image_from_docx(self, output_dir="./images"):
        """
        Извлекает изображения из .docx файла и возвращает список путей к сохранённым изображениям.
        :param output_dir: Директория для сохранения изображений (по умолчанию текущая директория)
        :return: Список путей к изображениям
        """
        try:
            # Проверяем, существует ли директория, и создаём её, если она отсутствует
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # Извлекаем изображения с помощью docx2txt
            # Изображения автоматически сохраняются в output_dir с именами image{число}.{расширение}
            docx2txt.process(self.path, output_dir)

            # Получаем список всех файлов в директории
            files_in_dir = os.listdir(output_dir)
            image_paths = []  # Создаём список для хранения путей к изображениям

            # Определяем допустимые расширения изображений
            valid_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
            for file_name in files_in_dir:
                # Проверяем, начинается ли имя файла с 'image' и имеет ли он допустимое расширение
                if file_name.startswith('image') and any(file_name.lower().endswith(ext) for ext in valid_extensions):
                    image_paths.append(os.path.join(output_dir, file_name))  # Добавляем полный путь к изображению

            # Сортируем пути для упорядоченного вывода (например, image1.png, image2.jpg)
            image_paths.sort()

            # Выводим информацию о найденных изображениях
            if image_paths:
                print(f"Найдено {len(image_paths)} изображений:")
                for path in image_paths:
                    print(f" - {path}")
            else:
                print("Изображения не найдены в документе")

            return image_paths  # Возвращаем список путей к изображениям

        except Exception as e:
            print(f"Ошибка при извлечении изображений: {e}")
            return []  # Возвращаем пустой список в случае ошибки

    def make_description(self):
        """
        Формирует текстовое описание документа, включая текст, таблицы и описания изображений.
        :return: Строку с описанием содержимого документа и изображений
        """
        text = self.text_extract()  # Извлекаем текст и таблицы из документа
        images_paths = self.det_image_from_docx()  # Извлекаем пути к изображениям
        describe_images = [ImageExtractions(path).gpt_describe() for path in images_paths]  # Получаем описания изображений с помощью ImageExtractions

        docx_data = ''  # Создаём пустую строку для хранения результата
        count = 1  # Счётчик для нумерации описаний изображений
        for i in text.items():  # Перебираем элементы словаря с текстом и таблицами
            docx_data += f'{i}\n'  # Добавляем текст или таблицу в результат
        for i in describe_images:  # Перебираем описания изображений
            docx_data += f'Описание картинки{count}: {i}\n'  # Добавляем описание изображения с номером
            count += 1  # Увеличиваем счётчик

        return docx_data  # Возвращаем итоговую строку с описанием